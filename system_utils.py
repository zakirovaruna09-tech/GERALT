# ══════════════════════════════════════════════════════════════════════
#  GERALT v6.0 — system_utils.py
#  Системные функции: статус ПК, приложения, email, Steam, docx
# ══════════════════════════════════════════════════════════════════════

import os, subprocess, webbrowser, smtplib, psutil, platform, shutil
from datetime import datetime
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from config import GMAIL, GMAIL_APP_PASS, STEAM_GAMES, JARVIS_DIR

# ══════════════════════════════════════════════════════════════════════
# СТАТУС ПК
# ══════════════════════════════════════════════════════════════════════
def get_pc_status() -> str:
    cpu     = psutil.cpu_percent(interval=1)
    ram     = psutil.virtual_memory()
    disk_path = "C:\\" if platform.system() == "Windows" else "/"
    disk    = psutil.disk_usage(disk_path)
    boot    = datetime.fromtimestamp(psutil.boot_time())
    uptime  = datetime.now() - boot
    hours   = int(uptime.total_seconds() // 3600)
    minutes = int((uptime.total_seconds() % 3600) // 60)

    temp_str = "н/д"
    try:
        temps = psutil.sensors_temperatures()
        if temps:
            for name, entries in temps.items():
                if entries:
                    temp_str = f"{entries[0].current:.0f}°C"
                    break
    except Exception:
        pass

    lines = [
        f"── СТАТУС СИСТЕМЫ ───────────────────",
        f"  CPU:      {cpu:.1f}%",
        f"  RAM:      {ram.percent:.1f}%  ({ram.used // 1024**3}/{ram.total // 1024**3} ГБ)",
        f"  Диск C:   {disk.percent:.1f}%  ({disk.used // 1024**3}/{disk.total // 1024**3} ГБ)",
        f"  Темп:     {temp_str}",
        f"  Аптайм:   {hours}ч {minutes}м",
        f"  ОС:       {platform.system()} {platform.release()}",
        f"─────────────────────────────────────",
    ]
    return "\n".join(lines)


def get_top_processes(n=5) -> str:
    procs = sorted(psutil.process_iter(["name", "cpu_percent", "memory_percent"]),
                   key=lambda p: p.info["cpu_percent"] or 0, reverse=True)[:n]
    lines = ["── ТОП ПРОЦЕССОВ ────────────────────"]
    for p in procs:
        lines.append(f"  {p.info['name'][:25]:<25}  CPU:{p.info['cpu_percent']:>5.1f}%  RAM:{p.info['memory_percent']:>4.1f}%")
    lines.append("─────────────────────────────────────")
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════
# ОТКРЫТИЕ ПРИЛОЖЕНИЙ
# ══════════════════════════════════════════════════════════════════════
APP_MAP = {
    "блокнот":       "notepad.exe",
    "notepad":       "notepad.exe",
    "браузер":       "chrome.exe",
    "chrome":        "chrome.exe",
    "firefox":       "firefox.exe",
    "калькулятор":   "calc.exe",
    "проводник":     "explorer.exe",
    "explorer":      "explorer.exe",
    "диспетчер":     "taskmgr.exe",
    "taskmgr":       "taskmgr.exe",
    "cmd":           "cmd.exe",
    "консоль":       "cmd.exe",
    "steam":         r"C:\Program Files (x86)\Steam\steam.exe",
    "telegram":      r"%APPDATA%\Telegram Desktop\Telegram.exe",
    "телеграм":      r"%APPDATA%\Telegram Desktop\Telegram.exe",
    "discord":       r"%APPDATA%\discord\Discord.exe",
    "vs code":       r"%LOCALAPPDATA%\Programs\Microsoft VS Code\Code.exe",
    "vscode":        r"%LOCALAPPDATA%\Programs\Microsoft VS Code\Code.exe",
    "spotify":       r"%APPDATA%\Spotify\Spotify.exe",
    "paint":         "mspaint.exe",
    "word":          "winword.exe",
    "excel":         "excel.exe",
    "яндекс музыка": "https://music.yandex.ru",
    "яндекс":        "https://music.yandex.ru",
    "музыка":        "https://music.yandex.ru",
    "youtube":       "https://youtube.com",
    "ютуб":          "https://youtube.com",
}

# Папки Desktop для поиска ярлыков
_DESKTOP_PATHS = [
    os.path.expandvars(r"%USERPROFILE%\Desktop"),
    os.path.expandvars(r"%PUBLIC%\Desktop"),
]


def _find_lnk(name: str) -> str | None:
    """Ищет ярлык .lnk на рабочем столе по имени."""
    for desktop in _DESKTOP_PATHS:
        if not os.path.isdir(desktop):
            continue
        for fname in os.listdir(desktop):
            if fname.lower().endswith(".lnk") and name.lower() in fname.lower():
                return os.path.join(desktop, fname)
    return None


def open_app(name: str) -> str:
    key = name.lower().strip()
    exe = APP_MAP.get(key)

    # URL → открываем в браузере
    if exe and exe.startswith("http"):
        webbrowser.open(exe)
        return f"Открываю {name} в браузере, сэр."

    # Если не нашли в словаре — пробуем имя напрямую
    if not exe:
        exe = name

    # ВАЖНО: раскрываем переменные окружения (%APPDATA% и т.д.)
    path = os.path.expandvars(exe)

    # 1. Пробуем os.startfile если файл существует
    if os.path.exists(path):
        try:
            os.startfile(path)
            return f"Открываю {name}, сэр."
        except Exception as e:
            return f"Сэр, не удалось открыть {name}: {e}"

    # 2. Ищем ярлык на рабочем столе
    lnk = _find_lnk(key)
    if lnk:
        try:
            os.startfile(lnk)
            return f"Открываю {name}, сэр."
        except Exception as e:
            return f"Сэр, не удалось открыть ярлык {name}: {e}"

    # 3. Пробуем как системную команду через subprocess (notepad.exe, calc.exe и т.д.)
    try:
        subprocess.Popen(path, shell=True,
                         creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, 'CREATE_NO_WINDOW') else 0)
        return f"Открываю {name}, сэр."
    except Exception as e:
        return f"Сэр, не удалось открыть {name}: {e}"


def close_app(name: str) -> str:
    killed = 0
    for proc in psutil.process_iter(["name"]):
        try:
            if name.lower() in proc.info["name"].lower():
                proc.kill()
                killed += 1
        except (psutil.NoSuchProcess, psutil.AccessDenied):
            pass
    if killed:
        return f"Завершил {killed} процесс(ов) {name}, сэр."
    return f"Сэр, процесс {name} не найден."


# ══════════════════════════════════════════════════════════════════════
# STEAM
# ══════════════════════════════════════════════════════════════════════
def launch_steam_game(name: str) -> str:
    key = name.lower().strip()
    for game_name, url in STEAM_GAMES.items():
        if game_name in key or key in game_name:
            webbrowser.open(url)
            return f"Запускаю {game_name} через Steam, сэр."
    return f"Сэр, игра '{name}' не найдена в списке Steam."


def download_steam_game(app_id: str) -> str:
    try:
        webbrowser.open(f"steam://install/{app_id}")
        return f"Открыта установка игры {app_id} через Steam, сэр."
    except Exception as e:
        return f"Сэр, ошибка: {e}"


# ══════════════════════════════════════════════════════════════════════
# EMAIL
# ══════════════════════════════════════════════════════════════════════
def send_email(to: str, subject: str, body: str, attachment_path: str = None) -> str:
    try:
        msg = MIMEMultipart()
        msg["From"]    = GMAIL
        msg["To"]      = to
        msg["Subject"] = subject
        msg.attach(MIMEText(body, "plain", "utf-8"))

        if attachment_path and os.path.exists(attachment_path):
            with open(attachment_path, "rb") as f:
                part = MIMEBase("application", "octet-stream")
                part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition",
                            f"attachment; filename={os.path.basename(attachment_path)}")
            msg.attach(part)

        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(GMAIL, GMAIL_APP_PASS)
            server.sendmail(GMAIL, to, msg.as_string())
        return f"Письмо отправлено на {to}, сэр."
    except smtplib.SMTPAuthenticationError:
        return "Сэр, ошибка авторизации Gmail. Проверьте GMAIL и GMAIL_APP_PASS в config.py."
    except Exception as e:
        return f"Сэр, ошибка отправки: {e}"


# ══════════════════════════════════════════════════════════════════════
# WORD ДОКУМЕНТЫ
# ══════════════════════════════════════════════════════════════════════
def create_docx(title: str, content: str) -> str | None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        h = doc.add_heading(title, 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.runs[0]
        run.font.color.rgb = RGBColor(0x7c, 0x3a, 0xed)

        doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        doc.add_paragraph()

        for line in content.split("\n"):
            doc.add_paragraph(line)

        os.makedirs(JARVIS_DIR, exist_ok=True)
        fname = f"{title}_{datetime.now().strftime('%d%m%Y_%H%M%S')}.docx"
        fpath = os.path.join(JARVIS_DIR, fname)
        doc.save(fpath)
        return fpath
    except ImportError:
        return None
    except Exception:
        return None


# ══════════════════════════════════════════════════════════════════════
# СКРИНШОТ
# ══════════════════════════════════════════════════════════════════════
def take_screenshot() -> str | None:
    try:
        import pyautogui
        os.makedirs(JARVIS_DIR, exist_ok=True)
        fname = f"screen_{datetime.now().strftime('%d%m%Y_%H%M%S')}.png"
        fpath = os.path.join(JARVIS_DIR, fname)
        pyautogui.screenshot(fpath)
        return fpath
    except Exception:
        return None


# ══════════════════════════════════════════════════════════════════════
# ПОГОДА
# ══════════════════════════════════════════════════════════════════════
def get_weather(city: str = "Tashkent") -> str:
    try:
        import requests as req
        r = req.get(f"https://wttr.in/{city}?format=3&lang=ru", timeout=5)
        if r.status_code == 200:
            return r.text.strip()
        return "Сэр, погода недоступна."
    except Exception:
        return "Сэр, погода недоступна."
